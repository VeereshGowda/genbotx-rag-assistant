"""
Content Manager for GenBotX
Handles file uploads, webpage scraping, and duplicate detection.
"""

import os
import hashlib
import requests
from pathlib import Path
from typing import List, Dict, Any, Optional, Set
from urllib.parse import urlparse, urljoin
import mimetypes
import time
from datetime import datetime

import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from langchain.schema import Document
from loguru import logger

class ContentManager:
    def __init__(self, documents_folder: str = "documents", upload_folder: str = "uploads"):
        self.documents_folder = Path(documents_folder)
        self.upload_folder = Path(upload_folder)
        self.documents_folder.mkdir(exist_ok=True)
        self.upload_folder.mkdir(exist_ok=True)
        
        # Track processed content to avoid duplicates
        self.content_hashes_file = self.documents_folder / "content_hashes.json"
        self.content_hashes = self._load_content_hashes()
        
        # Configure logging
        logger.add("logs/content_manager.log", rotation="10 MB", level="INFO")
        
    def _load_content_hashes(self) -> Dict[str, Dict[str, Any]]:
        """Load existing content hashes to detect duplicates"""
        try:
            if self.content_hashes_file.exists():
                import json
                with open(self.content_hashes_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            logger.warning(f"Could not load content hashes: {e}")
        return {}
    
    def _save_content_hashes(self):
        """Save content hashes to file"""
        try:
            import json
            with open(self.content_hashes_file, 'w', encoding='utf-8') as f:
                json.dump(self.content_hashes, f, indent=2)
        except Exception as e:
            logger.error(f"Could not save content hashes: {e}")
    
    def _get_content_hash(self, content: str) -> str:
        """Generate hash for content to detect duplicates"""
        return hashlib.md5(content.encode('utf-8')).hexdigest()
    
    def _is_duplicate_content(self, content: str, source: str) -> bool:
        """Check if content is already processed"""
        content_hash = self._get_content_hash(content)
        
        # Check if hash exists
        if content_hash in self.content_hashes:
            existing_info = self.content_hashes[content_hash]
            logger.info(f"Duplicate content detected. Original source: {existing_info['source']}, New source: {source}")
            return True
        
        return False
    
    def _register_content(self, content: str, source: str, content_type: str):
        """Register processed content to prevent duplicates"""
        content_hash = self._get_content_hash(content)
        self.content_hashes[content_hash] = {
            "source": source,
            "content_type": content_type,
            "processed_at": datetime.now().isoformat(),
            "content_length": len(content)
        }
        self._save_content_hashes()
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            return ""
    
    def clean_web_content(self, html_content: str, url: str) -> str:
        """Clean and extract text from HTML content"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove unwanted elements
            for element in soup.find_all(['script', 'style', 'nav', 'header', 'footer', 
                                        'aside', 'advertisement', 'ads']):
                element.decompose()
            
            # Remove elements with common ad/navigation classes
            for element in soup.find_all(class_=['nav', 'navigation', 'menu', 'sidebar', 
                                                'advertisement', 'ad', 'ads', 'footer', 
                                                'header', 'social', 'share']):
                element.decompose()
            
            # Extract main content - try different selectors
            main_content = None
            content_selectors = [
                'main', 'article', '[role="main"]', '.content', '.main-content',
                '.post-content', '.entry-content', '#content', '#main'
            ]
            
            for selector in content_selectors:
                main_content = soup.select_one(selector)
                if main_content:
                    break
            
            # If no main content found, use body
            if not main_content:
                main_content = soup.find('body')
            
            if not main_content:
                main_content = soup
            
            # Extract text
            text = main_content.get_text(separator='\n', strip=True)
            
            # Clean up text
            lines = text.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line = line.strip()
                if line and len(line) > 10:  # Filter out very short lines
                    cleaned_lines.append(line)
            
            cleaned_text = '\n'.join(cleaned_lines)
            
            # Remove excessive whitespace
            import re
            cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
            cleaned_text = re.sub(r' {2,}', ' ', cleaned_text)
            
            return cleaned_text
            
        except Exception as e:
            logger.error(f"Error cleaning web content from {url}: {e}")
            return ""
    
    def scrape_webpage(self, url: str) -> Optional[Dict[str, str]]:
        """Scrape and clean content from a webpage"""
        try:
            logger.info(f"Scraping webpage: {url}")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Get title
            soup = BeautifulSoup(response.content, 'html.parser')
            title_element = soup.find('title')
            title = title_element.get_text().strip() if title_element else urlparse(url).netloc
            
            # Clean content
            cleaned_content = self.clean_web_content(response.text, url)
            
            if not cleaned_content:
                logger.warning(f"No content extracted from {url}")
                return None
            
            return {
                'title': title,
                'content': cleaned_content,
                'url': url,
                'scraped_at': datetime.now().isoformat()
            }
            
        except requests.RequestException as e:
            logger.error(f"Request error scraping {url}: {e}")
        except Exception as e:
            logger.error(f"Error scraping webpage {url}: {e}")
        
        return None
    
    def process_uploaded_file(self, file_path: Path, original_filename: str) -> Optional[Document]:
        """Process an uploaded file and return a Document if not duplicate"""
        try:
            logger.info(f"Processing uploaded file: {original_filename}")
            
            # Extract text based on file type
            file_extension = Path(original_filename).suffix.lower()
            
            if file_extension == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                text = self.extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return None
            
            if not text.strip():
                logger.warning(f"No text extracted from {original_filename}")
                return None
            
            # Check for duplicates
            if self._is_duplicate_content(text, original_filename):
                logger.info(f"Skipping duplicate file: {original_filename}")
                return None
            
            # Register content
            self._register_content(text, original_filename, f"uploaded_file_{file_extension}")
            
            # Create metadata
            metadata = {
                "source": original_filename,
                "filename": original_filename,
                "file_type": file_extension,
                "content_type": "uploaded_file",
                "processed_at": datetime.now().isoformat(),
                "file_size": file_path.stat().st_size if file_path.exists() else 0
            }
            
            # Save processed file to documents folder
            safe_filename = self._get_safe_filename(original_filename)
            saved_path = self.documents_folder / f"uploaded_{safe_filename}"
            
            try:
                if file_extension == '.txt':
                    with open(saved_path, 'w', encoding='utf-8') as f:
                        f.write(f"Title: {original_filename}\n")
                        f.write(f"Source: Uploaded file\n")
                        f.write("=" * 50 + "\n\n")
                        f.write(text)
                else:
                    # For PDF/DOCX, save as text file with metadata
                    with open(saved_path.with_suffix('.txt'), 'w', encoding='utf-8') as f:
                        f.write(f"Title: {original_filename}\n")
                        f.write(f"Source: Uploaded file ({file_extension})\n")
                        f.write(f"Processed: {datetime.now().isoformat()}\n")
                        f.write("=" * 50 + "\n\n")
                        f.write(text)
                
                metadata["saved_path"] = str(saved_path)
                
            except Exception as e:
                logger.warning(f"Could not save processed file: {e}")
            
            # Create Document object
            document = Document(
                page_content=text,
                metadata=metadata
            )
            
            logger.info(f"Successfully processed uploaded file: {original_filename}")
            return document
            
        except Exception as e:
            logger.error(f"Error processing uploaded file {original_filename}: {e}")
            return None
    
    def process_webpage_url(self, url: str) -> Optional[Document]:
        """Process a webpage URL and return a Document if not duplicate"""
        try:
            # Scrape webpage
            scraped_data = self.scrape_webpage(url)
            if not scraped_data:
                return None
            
            content = scraped_data['content']
            title = scraped_data['title']
            
            # Check for duplicates
            if self._is_duplicate_content(content, url):
                logger.info(f"Skipping duplicate webpage: {url}")
                return None
            
            # Register content
            self._register_content(content, url, "scraped_webpage")
            
            # Create metadata
            metadata = {
                "source": url,
                "title": title,
                "content_type": "scraped_webpage",
                "scraped_at": scraped_data['scraped_at'],
                "domain": urlparse(url).netloc
            }
            
            # Save scraped content
            try:
                safe_filename = self._get_safe_filename(f"webpage_{urlparse(url).netloc}_{title}")
                saved_path = self.documents_folder / f"{safe_filename}.txt"
                
                with open(saved_path, 'w', encoding='utf-8') as f:
                    f.write(f"Title: {title}\n")
                    f.write(f"URL: {url}\n")
                    f.write(f"Scraped: {scraped_data['scraped_at']}\n")
                    f.write("=" * 50 + "\n\n")
                    f.write(content)
                
                metadata["saved_path"] = str(saved_path)
                
            except Exception as e:
                logger.warning(f"Could not save scraped content: {e}")
            
            # Create Document object
            document = Document(
                page_content=content,
                metadata=metadata
            )
            
            logger.info(f"Successfully processed webpage: {url}")
            return document
            
        except Exception as e:
            logger.error(f"Error processing webpage {url}: {e}")
            return None
    
    def process_multiple_urls(self, urls: List[str], delay: float = 1.0) -> List[Document]:
        """Process multiple webpage URLs with delay between requests"""
        documents = []
        
        for i, url in enumerate(urls):
            try:
                document = self.process_webpage_url(url)
                if document:
                    documents.append(document)
                
                # Add delay between requests to be respectful
                if i < len(urls) - 1:
                    time.sleep(delay)
                    
            except Exception as e:
                logger.error(f"Error processing URL {url}: {e}")
                continue
        
        logger.info(f"Processed {len(documents)} unique webpages out of {len(urls)} URLs")
        return documents
    
    def _get_safe_filename(self, filename: str) -> str:
        """Generate safe filename"""
        import re
        # Remove invalid characters
        safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        safe_filename = safe_filename.replace(' ', '_')
        return safe_filename[:100]  # Limit length
    
    def get_content_stats(self) -> Dict[str, Any]:
        """Get statistics about processed content"""
        return {
            "total_unique_content": len(self.content_hashes),
            "content_types": {
                content_type: sum(1 for info in self.content_hashes.values() 
                                if info.get("content_type", "").startswith(content_type))
                for content_type in ["uploaded_file", "scraped_webpage", "wikipedia"]
            },
            "total_content_length": sum(info.get("content_length", 0) 
                                      for info in self.content_hashes.values()),
            "upload_folder": str(self.upload_folder),
            "documents_folder": str(self.documents_folder)
        }
    
    def clear_duplicates_cache(self):
        """Clear the duplicates cache (use with caution)"""
        self.content_hashes = {}
        self._save_content_hashes()
        logger.info("Cleared content duplicates cache")

def main():
    """Test the content manager"""
    Path("logs").mkdir(exist_ok=True)
    
    manager = ContentManager()
    
    # Test stats
    stats = manager.get_content_stats()
    print("Content Manager Stats:")
    import json
    print(json.dumps(stats, indent=2))
    
    # Test URL processing
    test_urls = ["https://en.wikipedia.org/wiki/Python_(programming_language)"]
    documents = manager.process_multiple_urls(test_urls)
    print(f"Processed {len(documents)} documents from URLs")

if __name__ == "__main__":
    main()
