"""
Wikipedia Scraper for GenBotX
Scrapes Wikipedia URLs, cleans the data, and saves them as documents for RAG processing.
"""

import os
import re
import wikipediaapi
import requests
from bs4 import BeautifulSoup
from pathlib import Path
from urllib.parse import urlparse
from loguru import logger
from docx import Document as DocxDocument
from typing import List, Dict, Optional
from .config import get_config

class WikipediaScraper:
    def __init__(self, documents_folder: Optional[str] = None):
        config = get_config()
        
        self.documents_folder = Path(documents_folder or config.get("directories.documents"))
        self.documents_folder.mkdir(exist_ok=True)
        
        # Configure loguru logger
        log_config = config.get("logging")
        logger.add(
            f"{config.get('directories.logs')}/wikipedia_scraper.log", 
            rotation=log_config.get("rotation", "10 MB"),
            level=log_config.get("level", "INFO"),
            format=log_config.get("format")
        )
        
        # Get Wikipedia configuration
        wiki_config = config.get("wikipedia")
        self.language = wiki_config.get("language", "en")
        self.user_agent = wiki_config.get("user_agent")
        
        # Initialize Wikipedia API
        self.wiki_wiki = wikipediaapi.Wikipedia(
            language=self.language,
            extract_format=wikipediaapi.ExtractFormat.WIKI,
            user_agent=self.user_agent
        )
        
    def extract_title_from_url(self, url: str) -> str:
        """Extract Wikipedia article title from URL"""
        try:
            # Parse URL and extract the page title
            parsed_url = urlparse(url)
            title = parsed_url.path.split('/')[-1]
            title = title.replace('_', ' ')
            return title
        except Exception as e:
            logger.error(f"Error extracting title from URL {url}: {e}")
            return "Unknown"
    
    def clean_text(self, text: str) -> str:
        """Clean and format text content"""
        # Remove multiple whitespaces and newlines
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Remove citation markers like [1], [citation needed], etc.
        text = re.sub(r'\[.*?\]', '', text)
        
        # Remove extra parenthetical references
        text = re.sub(r'\(.*?pronunciation.*?\)', '', text, flags=re.IGNORECASE)
        
        return text.strip()
    
    def scrape_wikipedia_page(self, url: str) -> Dict[str, str]:
        """Scrape a single Wikipedia page and return cleaned content"""
        try:
            title = self.extract_title_from_url(url)
            logger.info(f"Scraping Wikipedia page: {title}")
            
            # Try using wikipediaapi library first
            try:
                page = self.wiki_wiki.page(title)
                if page.exists():
                    content = {
                        'title': page.title,
                        'summary': self.clean_text(page.summary),
                        'content': self.clean_text(page.text),
                        'url': url
                    }
                    logger.info(f"Successfully scraped {title} using wikipediaapi library")
                    return content
                else:
                    logger.warning(f"Page does not exist: {title}")
                    raise Exception(f"Page {title} does not exist")
            except Exception as wiki_error:
                logger.warning(f"Wikipediaapi library failed for {title}: {wiki_error}")
                
                # Fallback to BeautifulSoup scraping
                response = requests.get(url)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Extract title
                title_element = soup.find('h1', {'class': 'firstHeading'})
                title = title_element.text if title_element else self.extract_title_from_url(url)
                
                # Extract main content
                content_div = soup.find('div', {'id': 'mw-content-text'})
                if not content_div:
                    raise Exception("Could not find main content div")
                
                # Remove unwanted elements
                for element in content_div.find_all(['table', 'div', 'span'], 
                                                  class_=['navbox', 'infobox', 'citation', 'reference']):
                    element.decompose()
                
                # Extract text from paragraphs
                paragraphs = content_div.find_all('p')
                content_text = '\n'.join([p.get_text() for p in paragraphs if p.get_text().strip()])
                
                content = {
                    'title': title,
                    'summary': content_text[:500] + "..." if len(content_text) > 500 else content_text,
                    'content': self.clean_text(content_text),
                    'url': url
                }
                
                logger.info(f"Successfully scraped {title} using BeautifulSoup")
                return content
                
        except Exception as e:
            logger.error(f"Error scraping {url}: {e}")
            return None
    
    def save_as_docx(self, content: Dict[str, str], filename: str):
        """Save content as Word document"""
        doc = DocxDocument()
        
        # Add title
        doc.add_heading(content['title'], 0)
        
        # Add URL
        doc.add_paragraph(f"Source: {content['url']}")
        doc.add_paragraph("")  # Empty line
        
        # Add summary
        if content['summary']:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(content['summary'])
            doc.add_paragraph("")  # Empty line
        
        # Add main content
        doc.add_heading("Content", level=1)
        
        # Split content into paragraphs for better formatting
        paragraphs = content['content'].split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save document
        filepath = self.documents_folder / filename
        doc.save(str(filepath))
        logger.info(f"Saved document: {filepath}")
    
    def save_as_txt(self, content: Dict[str, str], filename: str):
        """Save content as text file"""
        filepath = self.documents_folder / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"Title: {content['title']}\n")
            f.write(f"Source: {content['url']}\n")
            f.write("=" * 50 + "\n\n")
            
            if content['summary']:
                f.write("SUMMARY:\n")
                f.write(content['summary'] + "\n\n")
            
            f.write("CONTENT:\n")
            f.write(content['content'])
        
        logger.info(f"Saved text file: {filepath}")
    
    def get_safe_filename(self, title: str, group: str = None) -> str:
        """Generate safe filename from title"""
        # Remove invalid characters
        safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)
        safe_title = safe_title.replace(' ', '_')
        
        if group:
            safe_title = f"{group}_{safe_title}"
        
        return safe_title[:100]  # Limit filename length
    
    def scrape_urls_from_file(self, urls_file: str = "documents/wikipedia_urls.txt"):
        """Scrape all URLs from the Wikipedia URLs file"""
        urls_path = Path(urls_file)
        if not urls_path.exists():
            logger.error(f"URLs file not found: {urls_file}")
            return
        
        with open(urls_path, 'r', encoding='utf-8') as f:
            urls = [line.strip() for line in f if line.strip() and not line.startswith('#')]
        
        logger.info(f"Found {len(urls)} URLs to scrape")
        
        # Group URLs by topic
        url_groups = {
            'transformers': [],
            'gupta_empire': [],
            'kuru_kingdom': [],
            'other': []
        }
        
        for url in urls:
            if 'transformer' in url.lower():
                url_groups['transformers'].append(url)
            elif 'gupta' in url.lower():
                url_groups['gupta_empire'].append(url)
            elif 'kuru' in url.lower():
                url_groups['kuru_kingdom'].append(url)
            else:
                url_groups['other'].append(url)
        
        for group, group_urls in url_groups.items():
            if not group_urls:
                continue
                
            logger.info(f"Processing {group} group with {len(group_urls)} URLs")
            
            for url in group_urls:
                try:
                    # Check if file already exists
                    title = self.extract_title_from_url(url)
                    safe_filename = self.get_safe_filename(title, group)
                    
                    txt_path = self.documents_folder / f"{safe_filename}.txt"
                    
                    if txt_path.exists():
                        logger.info(f"File already exists for {title}, skipping...")
                        continue
                    
                    # Scrape the page
                    content = self.scrape_wikipedia_page(url)
                    if content:
                        # Save as text file only
                        self.save_as_txt(content, f"{safe_filename}.txt")
                    
                except Exception as e:
                    logger.error(f"Error processing URL {url}: {e}")
                    continue
        
        logger.info("Wikipedia scraping completed!")

def main():
    """Main function to run the scraper"""
    # Create logs directory
    Path("logs").mkdir(exist_ok=True)
    
    scraper = WikipediaScraper()
    scraper.scrape_urls_from_file()

if __name__ == "__main__":
    main()
