"""
Document Processor for GenBotX
Handles document ingestion, processing, and chunking for the RAG system.
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
import PyPDF2
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from loguru import logger
import mimetypes
from .config import get_config

class DocumentProcessor:
    def __init__(self, documents_folder: Optional[str] = None):
        config = get_config()
        
        self.documents_folder = Path(documents_folder or config.get("directories.documents"))
        
        # Initialize text splitter with config values
        splitter_config = config.get("text_splitter")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=splitter_config.get("chunk_size", 1000),
            chunk_overlap=splitter_config.get("chunk_overlap", 200),
            length_function=len,
            separators=splitter_config.get("separators", ["\n\n", "\n", " ", ""])
        )
        
        # Configure loguru logger
        log_config = config.get("logging")
        logger.add(
            f"{config.get('directories.logs')}/document_processor.log", 
            rotation=log_config.get("rotation", "10 MB"),
            level=log_config.get("level", "INFO"),
            format=log_config.get("format")
        )
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            logger.info(f"Successfully extracted text from PDF: {file_path.name}")
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {file_path.name}")
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Successfully extracted text from TXT: {file_path.name}")
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            return ""
    
    def process_single_document(self, file_path: Path) -> List[Document]:
        """Process a single document and return chunks"""
        logger.info(f"Processing document: {file_path.name}")
        
        # Determine file type and extract text
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return []
        
        if not text:
            logger.warning(f"No text extracted from {file_path.name}")
            return []
        
        # Create metadata
        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "file_type": file_extension,
            "file_size": file_path.stat().st_size if file_path.exists() else 0
        }
        
        # Split text into chunks
        try:
            chunks = self.text_splitter.split_text(text)
            
            # Create Document objects with metadata
            documents = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata["chunk_id"] = i
                chunk_metadata["total_chunks"] = len(chunks)
                
                doc = Document(
                    page_content=chunk,
                    metadata=chunk_metadata
                )
                documents.append(doc)
            
            logger.info(f"Created {len(documents)} chunks from {file_path.name}")
            return documents
        
        except Exception as e:
            logger.error(f"Error chunking document {file_path.name}: {e}")
            return []
    
    def process_all_documents(self) -> List[Document]:
        """Process all documents in the documents folder"""
        if not self.documents_folder.exists():
            logger.error(f"Documents folder not found: {self.documents_folder}")
            return []
        
        all_documents = []
        supported_extensions = {'.pdf', '.docx', '.txt'}
        
        # Find all supported files
        files_to_process = []
        for file_path in self.documents_folder.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                # Skip the wikipedia_urls.txt file
                if file_path.name != 'wikipedia_urls.txt':
                    files_to_process.append(file_path)
        
        logger.info(f"Found {len(files_to_process)} documents to process")
        
        # Process each file
        for file_path in files_to_process:
            try:
                documents = self.process_single_document(file_path)
                all_documents.extend(documents)
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                continue
        
        logger.info(f"Total documents processed: {len(all_documents)} chunks from {len(files_to_process)} files")
        return all_documents
    
    def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        if not self.documents_folder.exists():
            return {"error": "Documents folder not found"}
        
        stats = {
            "total_files": 0,
            "file_types": {},
            "total_size": 0,
            "files": []
        }
        
        supported_extensions = {'.pdf', '.docx', '.txt'}
        
        for file_path in self.documents_folder.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                if file_path.name != 'wikipedia_urls.txt':
                    stats["total_files"] += 1
                    
                    file_type = file_path.suffix.lower()
                    stats["file_types"][file_type] = stats["file_types"].get(file_type, 0) + 1
                    
                    file_size = file_path.stat().st_size
                    stats["total_size"] += file_size
                    
                    stats["files"].append({
                        "name": file_path.name,
                        "type": file_type,
                        "size": file_size,
                        "path": str(file_path)
                    })
        
        return stats

def main():
    """Test the document processor"""
    # Create logs directory
    Path("logs").mkdir(exist_ok=True)
    
    processor = DocumentProcessor()
    
    # Get document statistics
    stats = processor.get_document_stats()
    print("Document Statistics:")
    print(f"Total files: {stats['total_files']}")
    print(f"File types: {stats['file_types']}")
    print(f"Total size: {stats['total_size']} bytes")
    
    # Process all documents
    documents = processor.process_all_documents()
    print(f"\nProcessed {len(documents)} document chunks")

if __name__ == "__main__":
    main()
